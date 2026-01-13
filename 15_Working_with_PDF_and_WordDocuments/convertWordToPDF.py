#  This script runs on Windows only, and you must have Word installed.

import os
import win32com.client
import collections.abc
collections.Sequence = collections.abc.Sequence
import docx
from docx.enum.text import WD_BREAK

wordFilename = os.path.abspath('your_word_document.docx')
pdfFilename = os.path.abspath('your_pdf_filename.pdf')

doc = docx.Document()
#  Code to write Word document goes here.
doc.add_paragraph('To write a program that produces PDFs with your own content, you '
                  'must use the docx module to create a Word document, then use the Pywin32 package’s'
                  'win32com.client module to convert it to a PDF. Replace the # Code to create Word '
                  'document goes here. comment with docx function calls to create your own content for '
                  'the PDF in a Word document.')
doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)
doc.add_paragraph('This may seem like a convoluted way to produce PDFs, but as it turns out, professional software '
                  'solutions are often just as complicated.')
doc.save(wordFilename)

wdFormatPDF = 17  # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()

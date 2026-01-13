# Reads a guest list and generates stylized Word invitations using
# Python-Docx, inserting page breaks between guests for batch printing.

import collections.abc
collections.Sequence = collections.abc.Sequence
import docx
from docx.enum.text import WD_BREAK

textFile = open(r"C:\Users\User\OneDrive\User\PycharmProjects\15_Working_with_PDF_and_WordDocuments\guests.txt")
listOfGuests = textFile.readlines()
print()

doc = docx.Document(r"C:\Users\User\Desktop\customInvitation.docx")


def write(index):
    doc.add_paragraph('It would be a pleasure to have the company of', 'Style1')
    doc.add_paragraph(listOfGuests[index].removesuffix('\n'), 'Style2')
    doc.add_paragraph('', 'Style1')
    doc.paragraphs[-1].add_run('at').underline = True
    doc.paragraphs[-1].add_run(' 11010 Memory Lane on the Evening of')
    doc.add_paragraph('April 1st', 'Style2')
    doc.add_paragraph('', 'Style1')
    doc.paragraphs[-1].add_run('at').underline = True
    doc.paragraphs[-1].add_run(' 7 o\'clock')
    doc.add_paragraph('')
    doc.paragraphs[-1].add_run('')
    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)


for guest in range(len(listOfGuests)):
    write(guest)

doc.save(r"C:\Users\User\Desktop\customInvitation.docx")
